"""Best-effort text extraction for uploaded documents (PDF / DOCX).

Deterministic given the same bytes; no LLM involved. Every function fails
CLOSED -- a corrupted, encrypted, or unparseable file returns None rather than
raising, so one bad upload never breaks the rest of a chat turn's context (the
caller, ContextWindowBuilder, treats None the same as "can't be read yet").

Scanned/image-only PDFs (no embedded text layer) also return None -- this is
text extraction, not OCR. That gap is a known follow-up (native PDF vision
input), not silently faked here.
"""

from __future__ import annotations

import io

from app.core.logger import logger

_MAX_PDF_PAGES = 200  # guard against pathological page counts, not the 25 MiB cap


def extract_pdf_text(content: bytes) -> str | None:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            # No password path here -- an encrypted upload stays unreadable
            # rather than guessing at credentials.
            return None
        pages = reader.pages[:_MAX_PDF_PAGES]
        text = "\n".join(page.extract_text() or "" for page in pages)
        return text.strip() or None
    except Exception as exc:  # noqa: BLE001 -- fail closed, never raise
        logger.warning("attachments: PDF text extraction failed",
                        extra={"stage": "extract_pdf", "error": str(exc)})
        return None


def extract_docx_text(content: bytes) -> str | None:
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text)
        text = "\n".join(parts)
        return text.strip() or None
    except Exception as exc:  # noqa: BLE001 -- fail closed, never raise
        logger.warning("attachments: DOCX text extraction failed",
                        extra={"stage": "extract_docx", "error": str(exc)})
        return None
